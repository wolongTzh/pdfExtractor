import docx

def readPara(path):
    file = docx.Document(path)
    for p in file.paragraphs:
        print(p.text)

def readTables(path):
    file = docx.Document(path)
    numTables = file.tables
    for table in numTables:
        row_count = len(table.rows)
        col_count = len(table.columns)
        for i in range(row_count):
            for j in range(col_count):
                print(table.cell(i, j).text)

if __name__ == '__main__':
    path = 'C:\\Users\\FEIFEI\\Desktop\\金融知识图谱项目\\test\\wps转换后报告.docx'
    readTables(path)